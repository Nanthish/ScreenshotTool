#!/usr/bin/env python3
"""
SnipIT - Windows Floating Screenshot Tool
A standalone application for taking screenshots with comments and timestamps,
then exporting them to a Word document.
"""

import sys
import os
import time
import threading
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from PIL import Image, ImageGrab, ImageTk, ImageDraw
import pythoncom
import win32clipboard
from docx import Document
from docx.shared import Inches, RGBColor
from datetime import datetime
import tempfile
import subprocess
import ctypes
from ctypes import wintypes
import win32gui
import win32con
import pyautogui
import win32com.client
import xml.etree.ElementTree as ET
import base64


class ScreenshotTool:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("SnipIT")
        self.root.geometry("140x130")
        self.root.attributes("-topmost", True)
        self.root.attributes("-alpha", 0.9)
        
        # Make window floating and borderless but keep it in taskbar
        self.root.overrideredirect(True)
        
        # Set window style to make it always on top and show in taskbar
        hwnd = self.root.winfo_id()
        ctypes.windll.user32.SetWindowLongW(hwnd, -20, 0x00000008 | 0x00000080 | 0x00000020 | 0x00040000)
        
        # Initialize data storage
        self.screenshots = []
        self.is_capturing = False
        self.add_comment_var = tk.BooleanVar(value=False)
        self.partial_screenshot_mode = False
        self.markup_mode = False
        self.select_window = None
        self.selection_overlay = None
        self.current_partial_image = None
        
        self.setup_ui()
        self.center_window()
        
        # Register global hotkeys (like Windows Snipping Tool)
        self.register_hotkeys()
        
    def register_hotkeys(self):
        """Register global hotkeys using Windows GetAsyncKeyState API"""
        try:
            # Virtual key codes
            VK_CONTROL = 0x11
            VK_MENU = 0x12  # Alt key
            VK_F = 0x46
            VK_P = 0x50
            
            # Function to check if key is pressed
            def is_key_pressed(vk):
                """Check if a virtual key is currently pressed"""
                return ctypes.windll.user32.GetAsyncKeyState(vk) & 0x8000
            
            def hotkey_listener():
                """Background thread that listens for hotkey presses"""
                last_f_pressed = False
                last_p_pressed = False
                
                while True:
                    try:
                        ctrl = is_key_pressed(VK_CONTROL)
                        alt = is_key_pressed(VK_MENU)
                        f_pressed = is_key_pressed(VK_F)
                        p_pressed = is_key_pressed(VK_P)
                        
                        # Ctrl+Alt+F
                        if ctrl and alt and f_pressed and not last_f_pressed:
                            print("✓ Ctrl+Alt+F detected - Taking screenshot...")
                            self.root.after(0, self.take_screenshot)
                            last_f_pressed = True
                        elif not f_pressed:
                            last_f_pressed = False
                        
                        # Ctrl+Alt+P
                        if ctrl and alt and p_pressed and not last_p_pressed:
                            print("✓ Ctrl+Alt+P detected - Starting partial capture...")
                            self.root.after(0, self.partial_capture)
                            last_p_pressed = True
                        elif not p_pressed:
                            last_p_pressed = False
                        
                        time.sleep(0.05)  # Poll every 50ms
                    except Exception as e:
                        print(f"Error in hotkey listener: {e}")
                        time.sleep(0.5)
            
            # Start listener thread
            listener_thread = threading.Thread(target=hotkey_listener, daemon=False)
            listener_thread.daemon = True
            listener_thread.start()
            
            print("✓ Global hotkeys ACTIVE: Ctrl+Alt+F (Full) and Ctrl+Alt+P (Partial)")
            print("✓ Works even when SnipIT is minimized or unfocused!")
            
        except Exception as e:
            print(f"✗ Failed to register hotkeys: {e}")
        
    def create_button_icons(self):
        """Create icon images for buttons using PIL"""
        # Screenshot icon (camera icon - 40x40)
        screenshot_icon = Image.new("RGBA", (40, 40), (255, 255, 255, 0))
        draw = ImageDraw.Draw(screenshot_icon)
        # Draw camera body
        draw.rectangle([8, 12, 32, 28], fill="#2196F3", outline="#1976D2", width=2)
        # Draw lens
        draw.ellipse([14, 16, 26, 28], fill="#64B5F6", outline="#1976D2", width=2)
        # Draw flash
        draw.rectangle([28, 8, 32, 12], fill="#FFC107", outline="#FFA000", width=1)
        self.screenshot_icon = ImageTk.PhotoImage(screenshot_icon)
        
        # End icon (document icon - 40x40)
        end_icon = Image.new("RGBA", (40, 40), (255, 255, 255, 0))
        draw = ImageDraw.Draw(end_icon)
        # Draw document
        draw.rectangle([10, 8, 30, 32], fill="#4CAF50", outline="#388E3C", width=2)
        # Draw lines on document
        draw.line([14, 14, 26, 14], fill="white", width=2)
        draw.line([14, 18, 26, 18], fill="white", width=2)
        draw.line([14, 22, 26, 22], fill="white", width=2)
        draw.line([14, 26, 22, 26], fill="white", width=2)
        self.end_icon = ImageTk.PhotoImage(end_icon)
        
        # Help icon (information i - rectangular blue - 32x32)
        help_icon = Image.new("RGBA", (32, 32), (255, 255, 255, 0))
        draw = ImageDraw.Draw(help_icon)
        # Draw rectangular blue background (same as X button but blue)
        draw.rectangle([4, 4, 28, 28], fill="#1976D2", outline="#1565C0", width=2)
        # Draw white "i" icon
        # Dot at top
        draw.ellipse([15, 9, 17, 11], fill="white")
        # Vertical line
        draw.line([16, 13, 16, 21], fill="white", width=3)
        self.help_icon = ImageTk.PhotoImage(help_icon)
        
        # Close icon (X mark - small rectangular box - 32x32)
        close_icon = Image.new("RGBA", (32, 32), (255, 255, 255, 0))
        draw = ImageDraw.Draw(close_icon)
        # Draw red rectangular background
        draw.rectangle([4, 4, 28, 28], fill="#E63946", outline="#C1121F", width=2)
        # Draw X mark with white lines
        draw.line([10, 10, 22, 22], fill="white", width=2)
        draw.line([22, 10, 10, 22], fill="white", width=2)
        self.close_icon = ImageTk.PhotoImage(close_icon)
        
        # Partial screenshot icon (crosshair icon - 40x40)
        partial_icon = Image.new("RGBA", (40, 40), (255, 255, 255, 0))
        draw = ImageDraw.Draw(partial_icon)
        # Draw crosshair
        draw.line([20, 5, 20, 35], fill="#FF9800", width=2)
        draw.line([5, 20, 35, 20], fill="#FF9800", width=2)
        # Draw circle in center
        draw.ellipse([17, 17, 23, 23], outline="#FF9800", width=2)
        self.partial_icon = ImageTk.PhotoImage(partial_icon)
        
        # Markup icon (pen icon - 40x40)
        markup_icon = Image.new("RGBA", (40, 40), (255, 255, 255, 0))
        draw = ImageDraw.Draw(markup_icon)
        # Draw pen
        draw.line([10, 30, 30, 10], fill="#E91E63", width=3)
        draw.ellipse([28, 8, 34, 14], fill="#E91E63")
        self.markup_icon = ImageTk.PhotoImage(markup_icon)
    
    def setup_ui(self):
        """Setup the floating UI with buttons"""
        self.create_button_icons()
        
        # Main frame with thin black border
        main_frame = tk.Frame(self.root, bg="#f0f0f0", padx=1, pady=1, relief=tk.SOLID, bd=1, highlightthickness=0, highlightbackground="black")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Bind drag to main frame
        main_frame.bind('<Button-1>', self.start_move)
        main_frame.bind('<B1-Motion>', self.do_move)
        
        # Top frame for title and close button
        top_frame = tk.Frame(main_frame, bg="#f0f0f0")
        top_frame.pack(fill=tk.X, pady=(0, 1))
        
        # Title label - make it draggable
        self.title_label = tk.Label(top_frame, text="SnipIT", 
                              font=("Arial", 7, "bold"), bg="#f0f0f0", fg="#333333", cursor="hand2")
        self.title_label.pack(side=tk.LEFT, expand=True)
        
        # Bind drag events to title label
        self.title_label.bind('<Button-1>', self.start_move)
        self.title_label.bind('<B1-Motion>', self.do_move)
        
        # Close button (red X) - small rectangular
        self.close_btn = tk.Button(top_frame, image=self.close_icon,
                                   command=self.close_tool,
                                   bg="#ffffff", relief=tk.FLAT, bd=0,
                                   cursor="hand2", width=24, height=24)
        self.close_btn.pack(side=tk.RIGHT, padx=(1, 0))
        
        # Button frame for main buttons
        button_frame = tk.Frame(main_frame, bg="#f0f0f0")
        button_frame.pack(fill=tk.X, pady=(0, 1))
        
        # Bind drag to button frame too
        button_frame.bind('<Button-1>', self.start_move)
        button_frame.bind('<B1-Motion>', self.do_move)
        
        # Screenshot button
        self.screenshot_btn = tk.Button(button_frame, image=self.screenshot_icon,
                                       command=self.take_screenshot, 
                                       bg="#ffffff", relief=tk.RAISED, bd=1,
                                       cursor="hand2")
        self.screenshot_btn.pack(side=tk.LEFT, padx=(0, 1))
        
        # Partial screenshot button (crosshair icon)
        self.partial_btn = tk.Button(button_frame, image=self.partial_icon,
                                     command=self.partial_capture,
                                     bg="#ffffff", relief=tk.RAISED, bd=1,
                                     cursor="hand2")
        self.partial_btn.pack(side=tk.LEFT, padx=(0, 1))
        
        # End button
        self.end_btn = tk.Button(button_frame, image=self.end_icon,
                                command=self.end_session,
                                bg="#ffffff", relief=tk.RAISED, bd=1,
                                cursor="hand2")
        self.end_btn.pack(side=tk.LEFT)
        
        # Checkbox frame for comments
        checkbox_frame = tk.Frame(main_frame, bg="#f0f0f0")
        checkbox_frame.pack(fill=tk.X, pady=(0, 1))
        
        self.comment_checkbox = tk.Checkbutton(checkbox_frame, text="Comment", 
                                               variable=self.add_comment_var,
                                               bg="#f0f0f0", font=("Arial", 6))
        self.comment_checkbox.pack(side=tk.LEFT)
        
        # Bind drag to checkbox frame
        checkbox_frame.bind('<Button-1>', self.start_move)
        checkbox_frame.bind('<B1-Motion>', self.do_move)
        
        # Bottom frame for help button (bottom right)
        bottom_frame = tk.Frame(main_frame, bg="#f0f0f0")
        bottom_frame.pack(fill=tk.X, pady=(0, 0))
        
        # Spacer to push help button to right
        spacer = tk.Frame(bottom_frame, bg="#f0f0f0")
        spacer.pack(side=tk.LEFT, expand=True)
        
        # Bind drag to bottom frame
        bottom_frame.bind('<Button-1>', self.start_move)
        bottom_frame.bind('<B1-Motion>', self.do_move)
        
        # Help button (blue) - rectangular, bottom right, same size as close button
        self.help_btn = tk.Button(bottom_frame, image=self.help_icon,
                                  command=self.show_help,
                                  bg="#ffffff", relief=tk.FLAT, bd=0,
                                  cursor="hand2", width=24, height=24)
        self.help_btn.pack(side=tk.RIGHT, padx=(0, 1))
        
    def partial_capture(self):
        """Start fullscreen overlay for region selection - Windows Snipping Tool style"""
        if self.is_capturing:
            return
            
        self.is_capturing = True
        self.partial_screenshot_mode = True
        
        try:
            # Keep main window handle
            main_hwnd = self.root.winfo_id()
            
            # Create fullscreen overlay window
            self.selection_overlay = tk.Tk()
            self.selection_overlay.attributes("-fullscreen", True)
            self.selection_overlay.attributes("-alpha", 0.3)  # Semi-transparent
            self.selection_overlay.attributes("-topmost", True)
            self.selection_overlay.config(bg="gray20")
            
            # Create canvas for drawing selection
            canvas = tk.Canvas(self.selection_overlay, bg="gray20", cursor="crosshair", 
                             highlightthickness=0, relief=tk.FLAT, bd=0)
            canvas.pack(fill=tk.BOTH, expand=True)
            
            # Get overlay window handle
            overlay_hwnd = self.selection_overlay.winfo_id()
            
            # Update window
            self.selection_overlay.update()
            
            # Use Windows API to make overlay transparent to mouse clicks but still visible
            HWND_TOPMOST = -1
            SWP_NOSIZE = 0x0001
            SWP_NOMOVE = 0x0002
            WS_EX_LAYERED = 0x00080000
            WS_EX_TRANSPARENT = 0x00000020
            
            # Set window position to topmost
            ctypes.windll.user32.SetWindowPos(overlay_hwnd, HWND_TOPMOST, 0, 0, 0, 0, SWP_NOSIZE | SWP_NOMOVE)
            
            # Small delay
            time.sleep(0.05)

            
            # Selection state
            self.selection_data = {
                "start_x": 0, "start_y": 0,
                "end_x": 0, "end_y": 0,
                "rect": None,
                "selecting": False
            }
            
            def on_mouse_down(event):
                """Handle mouse button down"""
                self.selection_data["start_x"] = event.x_root
                self.selection_data["start_y"] = event.y_root
                self.selection_data["end_x"] = event.x_root
                self.selection_data["end_y"] = event.y_root
                self.selection_data["selecting"] = True
                print(f"✓ Mouse down at ({event.x_root}, {event.y_root})")
                
                # Clear any previous rectangle
                if self.selection_data["rect"]:
                    canvas.delete(self.selection_data["rect"])
                    self.selection_data["rect"] = None
            
            def on_mouse_move(event):
                """Handle mouse motion while dragging"""
                if not self.selection_data["selecting"]:
                    return
                
                self.selection_data["end_x"] = event.x_root
                self.selection_data["end_y"] = event.y_root
                
                # Delete previous rectangle
                if self.selection_data["rect"]:
                    canvas.delete(self.selection_data["rect"])
                
                # Draw new selection rectangle
                x1 = min(self.selection_data["start_x"], self.selection_data["end_x"])
                y1 = min(self.selection_data["start_y"], self.selection_data["end_y"])
                x2 = max(self.selection_data["start_x"], self.selection_data["end_x"])
                y2 = max(self.selection_data["start_y"], self.selection_data["end_y"])
                
                # Draw rectangle in canvas coordinates
                canvas_x1 = x1 - self.selection_overlay.winfo_x()
                canvas_y1 = y1 - self.selection_overlay.winfo_y()
                canvas_x2 = x2 - self.selection_overlay.winfo_x()
                canvas_y2 = y2 - self.selection_overlay.winfo_y()
                
                self.selection_data["rect"] = canvas.create_rectangle(
                    canvas_x1, canvas_y1, canvas_x2, canvas_y2,
                    outline="white", width=2, fill="blue"
                )
                
                # Draw corner handles
                handle_size = 5
                for hx, hy in [(canvas_x1, canvas_y1), (canvas_x2, canvas_y1), 
                               (canvas_x1, canvas_y2), (canvas_x2, canvas_y2)]:
                    canvas.create_rectangle(
                        hx - handle_size, hy - handle_size,
                        hx + handle_size, hy + handle_size,
                        fill="white", outline="white"
                    )
            
            def on_mouse_up(event):
                """Handle mouse button release"""
                if not self.selection_data["selecting"]:
                    return
                
                self.selection_data["selecting"] = False
                self.selection_data["end_x"] = event.x_root
                self.selection_data["end_y"] = event.y_root
                
                # Validate selection size
                width = abs(self.selection_data["end_x"] - self.selection_data["start_x"])
                height = abs(self.selection_data["end_y"] - self.selection_data["start_y"])
                
                print(f"✓ Mouse released at ({event.x_root}, {event.y_root})")
                print(f"✓ Selection size: {width}x{height}")
                
                if width >= 10 and height >= 10:
                    # Valid selection - show countdown dialog for dropdown preparation
                    try:
                        print("✓ Destroying overlay...")
                        self.selection_overlay.destroy()
                        self.selection_overlay = None
                        
                        # Store selection coordinates
                        x1 = min(self.selection_data["start_x"], self.selection_data["end_x"])
                        y1 = min(self.selection_data["start_y"], self.selection_data["end_y"])
                        x2 = max(self.selection_data["start_x"], self.selection_data["end_x"])
                        y2 = max(self.selection_data["start_y"], self.selection_data["end_y"])
                        
                        print(f"✓ Selection bbox: ({x1}, {y1}, {x2}, {y2})")
                        
                        # Create a countdown window for dropdown preparation
                        countdown_window = tk.Toplevel(self.root)
                        countdown_window.attributes("-topmost", True)
                        countdown_window.title("Prepare Screenshot")
                        countdown_window.geometry("300x150")
                        countdown_window.resizable(False, False)
                        
                        # Center the countdown window
                        countdown_window.update_idletasks()
                        x = countdown_window.winfo_screenwidth() // 2 - 150
                        y = countdown_window.winfo_screenheight() // 2 - 75
                        countdown_window.geometry(f"300x150+{x}+{y}")
                        
                        # Label
                        label = tk.Label(countdown_window, text="Prepare capturing contents!\nGet your dropdown/menu ready.",
                                        font=("Arial", 10), fg="#333333")
                        label.pack(pady=10)
                        
                        # Countdown label
                        countdown_label = tk.Label(countdown_window, text="5", 
                                                  font=("Arial", 48, "bold"), fg="#2196F3")
                        countdown_label.pack(pady=10)
                        
                        # Instructions
                        instructions = tk.Label(countdown_window, text="Capturing in 5 seconds...",
                                              font=("Arial", 9), fg="#666666")
                        instructions.pack(pady=5)
                        
                        countdown_value = [5]
                        
                        def countdown():
                            countdown_value[0] -= 1
                            if countdown_value[0] > 0:
                                countdown_label.config(text=str(countdown_value[0]))
                                countdown_window.after(1000, countdown)
                            else:
                                # Time's up - hide countdown window and capture
                                countdown_window.withdraw()
                                countdown_window.update_idletasks()
                                self.root.update_idletasks()
                                time.sleep(0.3)  # Increased delay for reliability
                                perform_capture()
                        
                        def perform_capture():
                            """Perform the actual screenshot capture"""
                            try:
                                # Hide overlay and main window before capturing
                                if self.selection_overlay:
                                    self.selection_overlay.withdraw()
                                
                                self.root.withdraw()
                                self.root.update()
                                time.sleep(0.1)
                                
                                print(f"✓ Capturing bbox: ({x1}, {y1}, {x2}, {y2})")
                                self.current_partial_image = ImageGrab.grab(bbox=(x1, y1, x2, y2))
                                print(f"✓ Image captured: {self.current_partial_image.size}")
                                
                                # Show main window
                                self.root.deiconify()
                                
                                # Destroy overlay completely
                                try:
                                    if self.selection_overlay:
                                        self.selection_overlay.destroy()
                                        self.selection_overlay = None
                                except:
                                    pass
                                
                                # Open markup window
                                self.open_markup_window(self.current_partial_image)
                                print("✓ Markup window opened")
                                
                                self.is_capturing = False
                                self.partial_screenshot_mode = False
                            except Exception as e:
                                print(f"✗ Error in capture: {e}")
                                import traceback
                                traceback.print_exc()
                                self.root.deiconify()
                                self.is_capturing = False
                                self.partial_screenshot_mode = False
                        
                        # Start countdown
                        countdown_window.after(1000, countdown)
                        
                    except Exception as e:
                        print(f"✗ Error in capture: {e}")
                        import traceback
                        traceback.print_exc()
                        self.is_capturing = False
                        self.partial_screenshot_mode = False
                    
                else:
                    # Selection too small - show message but allow retry
                    print(f"✗ Selection too small: {width}x{height}")
                    # Restore main window before showing warning
                    try:
                        self.selection_overlay.destroy()
                        self.selection_overlay = None
                    except:
                        pass
                    # Restore main window visibility
                    self.root.deiconify()
                    messagebox.showwarning(
                        "Invalid Selection", 
                        "Selection must be at least 10x10 pixels.\nTry again."
                    )
                    # Hide again for retry
                    self.root.withdraw()
            
            def on_key_press(event):
                """Handle key presses"""
                if event.keysym == "Escape":
                    print("✓ Escape pressed - canceling")
                    try:
                        self.selection_overlay.destroy()
                        self.selection_overlay = None
                    except:
                        pass
                    # Restore main window
                    self.root.deiconify()
                    self.is_capturing = False
                    self.partial_screenshot_mode = False
            
            # Bind mouse events to both canvas and overlay window
            canvas.bind("<Button-1>", on_mouse_down)
            canvas.bind("<B1-Motion>", on_mouse_move)
            canvas.bind("<ButtonRelease-1>", on_mouse_up)
            canvas.bind("<KeyPress>", on_key_press)
            
            self.selection_overlay.bind("<Button-1>", on_mouse_down)
            self.selection_overlay.bind("<B1-Motion>", on_mouse_move)
            self.selection_overlay.bind("<ButtonRelease-1>", on_mouse_up)
            self.selection_overlay.bind("<KeyPress>", on_key_press)
            
            # Give focus to both
            self.selection_overlay.focus_set()
            canvas.focus_set()
            
            # Update to ensure window is fully rendered
            self.selection_overlay.update()
            
            print("✓ Partial capture started")
            print("✓ Drag to select region with crosshair cursor")
            print("✓ Press Escape to cancel")
            
        except Exception as e:
            print(f"✗ Error in partial_capture: {e}")
            import traceback
            traceback.print_exc()
            self.is_capturing = False
            self.partial_screenshot_mode = False
            messagebox.showerror("Error", f"Failed to start partial capture: {str(e)}")
        
    def center_window(self):
        """Position the floating window 45 pixels away from both corners"""
        try:
            screen_width = self.root.winfo_screenwidth()
            screen_height = self.root.winfo_screenheight()
            
            x = 45  # 45 pixels from left edge
            y = 45  # 45 pixels from top edge
            
            self.root.geometry(f"140x130+{x}+{y}")
        except tk.TclError:
            pass
        
    def start_move(self, event):
        """Start dragging the window"""
        # Only start move if not clicking on buttons
        if event.widget.__class__.__name__ != 'Button':
            self.x = event.x_root
            self.y = event.y_root
        
    def do_move(self, event):
        """Move the window while dragging"""
        # Only move if not on buttons
        if event.widget.__class__.__name__ != 'Button':
            deltax = event.x_root - self.x
            deltay = event.y_root - self.y
            x = self.root.winfo_x() + deltax
            y = self.root.winfo_y() + deltay
            self.root.geometry(f"+{x}+{y}")
            self.x = event.x_root
            self.y = event.y_root
        
    def take_screenshot(self):
        """Take a screenshot of the entire screen"""
        if self.is_capturing:
            return
            
        self.is_capturing = True
        
        try:
            # Hide the floating window
            self.root.withdraw()
            
            # Small delay to ensure window is hidden
            self.root.update()
            time.sleep(0.1)
            
            # Take screenshot
            screenshot = ImageGrab.grab()
            
            # Show window again
            self.root.deiconify()
            
            # Get comment only if checkbox is enabled
            comment = ""
            if self.add_comment_var.get():
                comment = self.get_comment()
                if comment is None:  # User clicked cancel
                    self.is_capturing = False
                    return
                    
            # Store screenshot data
            timestamp = datetime.now()
            self.screenshots.append({
                'image': screenshot,
                'comment': comment,
                'timestamp': timestamp
            })
            
            print("✓ Full screenshot captured")
                
        except Exception as e:
            self.root.deiconify()
            messagebox.showerror("Error", f"Failed to take screenshot: {str(e)}")
        finally:
            self.is_capturing = False
            
    def get_comment(self):
        """Get a comment from the user for the screenshot"""
        comment = simpledialog.askstring("Add Comment", 
                                       "Enter a comment for this screenshot:",
                                       parent=self.root)
        return comment if comment is not None else ""
        
    def take_partial_screenshot(self):
        """Allow user to select a region and take a partial screenshot"""
        if self.is_capturing:
            return
            
        self.is_capturing = True
        self.partial_screenshot_mode = True
        
        try:
            # Hide the floating window
            self.root.withdraw()
            self.root.update()
            time.sleep(0.1)
            
            # Create selection window
            self.create_selection_window()
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to take partial screenshot: {str(e)}")
        finally:
            self.is_capturing = False
            self.partial_screenshot_mode = False
            
    def smart_dropdown_capture(self):
        """Capture the focused window with auto-padding"""
        try:
            # Get focused window
            hwnd = win32gui.GetForegroundWindow()
            if not hwnd:
                messagebox.showerror("Error", "Could not get focused window.")
                return
            
            # Get window rectangle
            rect = win32gui.GetWindowRect(hwnd)
            x1, y1, x2, y2 = rect
            
            # Add 50px padding
            x1 = max(0, x1 - 50)
            y1 = max(0, y1 - 50)
            x2 = x2 + 50
            y2 = y2 + 50
            
            # Capture region
            self.current_partial_image = ImageGrab.grab(bbox=(x1, y1, x2, y2))
            
            # Open markup window
            self.open_markup_window(self.current_partial_image)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to capture window: {str(e)}")
    
    def open_markup_window(self, image):
        """Open a window to markup/annotate the screenshot"""
        markup_window = tk.Toplevel(self.root)
        markup_window.title("Partial Capture - Markup")
        markup_window.attributes("-topmost", True)
        
        # Calculate canvas size based on image dimensions (max 800px)
        img_width, img_height = image.size
        max_dim = 800
        
        if img_width > max_dim or img_height > max_dim:
            scale = min(max_dim / img_width, max_dim / img_height)
            display_width = int(img_width * scale)
            display_height = int(img_height * scale)
        else:
            display_width = img_width
            display_height = img_height
        
        # Set window geometry with proper dimensions
        window_width = max(display_width + 20, 500)
        window_height = display_height + 180
        markup_window.geometry(f"{window_width}x{window_height}")
        
        # Canvas for drawing (at top, with fixed size)
        canvas_frame = tk.Frame(markup_window)
        canvas_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        canvas = tk.Canvas(canvas_frame, bg="white", cursor="crosshair", highlightthickness=1, 
                          width=display_width, height=display_height)
        canvas.pack(fill=tk.BOTH, expand=True)
        
        # Display image on canvas (resized if needed)
        display_image = image.copy()
        if display_width != img_width or display_height != img_height:
            display_image = display_image.resize((display_width, display_height), Image.Resampling.LANCZOS)
        
        photo = ImageTk.PhotoImage(display_image)
        img_on_canvas = canvas.create_image(0, 0, image=photo, anchor="nw")
        canvas.image = photo
        canvas.photo = photo
        
        # Store original image reference for capture
        canvas.original_image = image
        canvas.display_image = display_image
        canvas.scale_factor = display_width / img_width if img_width > 0 else 1
        
        # Tool selection variables
        tool_var = tk.StringVar(value="rectangle")
        color_var = tk.StringVar(value="red")
        
        # Drawing state and markup storage
        drawing_data = {"drawing": False, "rect_id": None, "markups": []}
        
        def start_draw(event):
            drawing_data["drawing"] = True
            drawing_data["last_x"] = event.x
            drawing_data["last_y"] = event.y
            drawing_data["rect_id"] = None
        
        def on_mouse_drag(event):
            if not drawing_data["drawing"]:
                return
            
            tool = tool_var.get()
            color = color_var.get()
            
            # Delete previous rectangle preview if drawing rectangle or circle
            if tool in ["rectangle", "circle"] and drawing_data["rect_id"] is not None:
                canvas.delete(drawing_data["rect_id"])
            
            if tool == "rectangle":
                # Draw rectangle preview
                drawing_data["rect_id"] = canvas.create_rectangle(
                    drawing_data["last_x"], drawing_data["last_y"],
                    event.x, event.y,
                    outline=color, width=2
                )
            elif tool == "circle":
                # Draw circle preview
                drawing_data["rect_id"] = canvas.create_oval(
                    drawing_data["last_x"], drawing_data["last_y"],
                    event.x, event.y,
                    outline=color, width=2
                )
            elif tool == "draw":
                # Draw line
                canvas.create_line(
                    drawing_data["last_x"], drawing_data["last_y"],
                    event.x, event.y,
                    fill=color, width=2
                )
                drawing_data["last_x"] = event.x
                drawing_data["last_y"] = event.y
        
        def stop_draw(event):
            tool = tool_var.get()
            color = color_var.get()
            
            # Store markup data for later redrawing
            if drawing_data["drawing"]:
                # Scale coordinates back to original image size
                scale_factor = canvas.scale_factor
                x1 = int(drawing_data["last_x"] / scale_factor)
                y1 = int(drawing_data["last_y"] / scale_factor)
                x2 = int(event.x / scale_factor)
                y2 = int(event.y / scale_factor)
                
                if tool == "rectangle":
                    drawing_data["markups"].append({
                        "type": "rectangle",
                        "x1": x1,
                        "y1": y1,
                        "x2": x2,
                        "y2": y2,
                        "color": color
                    })
                elif tool == "circle":
                    drawing_data["markups"].append({
                        "type": "circle",
                        "x1": x1,
                        "y1": y1,
                        "x2": x2,
                        "y2": y2,
                        "color": color
                    })
                elif tool == "draw":
                    # Store draw strokes
                    drawing_data["markups"].append({
                        "type": "draw",
                        "x1": x1,
                        "y1": y1,
                        "x2": x2,
                        "y2": y2,
                        "color": color
                    })
            
            drawing_data["drawing"] = False
            drawing_data["rect_id"] = None
        
        canvas.bind("<Button-1>", start_draw)
        canvas.bind("<B1-Motion>", on_mouse_drag)
        canvas.bind("<ButtonRelease-1>", stop_draw)
        
        # Control frame at BOTTOM
        control_frame = tk.Frame(markup_window, bg="#f0f0f0")
        control_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=5, pady=5)
        
        # Tool selection (Rectangle, Circle, Draw) as radio buttons (vertical layout)
        tool_frame = tk.LabelFrame(control_frame, text="Markup Tool", bg="#f0f0f0", font=("Arial", 8), padx=2, pady=1)
        tool_frame.pack(side=tk.LEFT, padx=2, fill=tk.BOTH, expand=False)
        
        tk.Radiobutton(tool_frame, text="Rectangle", variable=tool_var, value="rectangle", 
                       bg="#f0f0f0", font=("Arial", 7)).pack(side=tk.TOP, padx=0, pady=1, anchor="w")
        tk.Radiobutton(tool_frame, text="Circle", variable=tool_var, value="circle", 
                       bg="#f0f0f0", font=("Arial", 7)).pack(side=tk.TOP, padx=0, pady=1, anchor="w")
        tk.Radiobutton(tool_frame, text="Draw", variable=tool_var, value="draw", 
                       bg="#f0f0f0", font=("Arial", 7)).pack(side=tk.TOP, padx=0, pady=1, anchor="w")
        
        # Color selection as radio buttons (vertical layout)
        color_frame = tk.LabelFrame(control_frame, text="Color", bg="#f0f0f0", font=("Arial", 8), padx=2, pady=1)
        color_frame.pack(side=tk.LEFT, padx=2, fill=tk.BOTH, expand=False)
        
        # Red color option with sample box
        red_container = tk.Frame(color_frame, bg="#f0f0f0")
        red_container.pack(side=tk.TOP, padx=0, pady=1, anchor="w")
        tk.Radiobutton(red_container, text="Red", variable=color_var, value="red", 
                       bg="#f0f0f0", font=("Arial", 7)).pack(side=tk.LEFT, padx=0)
        tk.Label(red_container, bg="red", width=2, height=1).pack(side=tk.LEFT, padx=1)
        
        # Yellow color option with sample box
        yellow_container = tk.Frame(color_frame, bg="#f0f0f0")
        yellow_container.pack(side=tk.TOP, padx=0, pady=1, anchor="w")
        tk.Radiobutton(yellow_container, text="Yellow", variable=color_var, value="yellow", 
                       bg="#f0f0f0", font=("Arial", 7)).pack(side=tk.LEFT, padx=0)
        tk.Label(yellow_container, bg="yellow", width=2, height=1).pack(side=tk.LEFT, padx=1)
        
        # Green color option with sample box
        green_container = tk.Frame(color_frame, bg="#f0f0f0")
        green_container.pack(side=tk.TOP, padx=0, pady=1, anchor="w")
        tk.Radiobutton(green_container, text="Green", variable=color_var, value="green", 
                       bg="#f0f0f0", font=("Arial", 7)).pack(side=tk.LEFT, padx=0)
        tk.Label(green_container, bg="green", width=2, height=1).pack(side=tk.LEFT, padx=1)
        
        # Save button - pass drawing_data to save_markup
        tk.Button(control_frame, text="Save", command=lambda: self.save_markup(markup_window, image, drawing_data),
                 bg="#4CAF50", fg="white", font=("Arial", 8)).pack(side=tk.LEFT, padx=5)
        
        # Clear button - clear all markups
        def clear_markups():
            drawing_data["markups"] = []
            # Redraw canvas with original image
            canvas.delete("all")
            canvas.create_image(0, 0, image=canvas.photo, anchor="nw")
        
        tk.Button(control_frame, text="Clear", command=clear_markups,
                 bg="#FF9800", fg="white", font=("Arial", 8)).pack(side=tk.LEFT, padx=2)
        
    def save_markup(self, window, original_image, drawing_data):
        """Save the marked up screenshot"""
        window.destroy()
        
        # Apply markups to the image using PIL
        marked_image = original_image.copy()
        
        if drawing_data["markups"]:
            draw = ImageDraw.Draw(marked_image)
            
            for markup in drawing_data["markups"]:
                if markup["type"] == "rectangle":
                    draw.rectangle(
                        [markup["x1"], markup["y1"], markup["x2"], markup["y2"]],
                        outline=markup["color"],
                        width=2
                    )
                elif markup["type"] == "circle":
                    draw.ellipse(
                        [markup["x1"], markup["y1"], markup["x2"], markup["y2"]],
                        outline=markup["color"],
                        width=2
                    )
                elif markup["type"] == "draw":
                    draw.line(
                        [markup["x1"], markup["y1"], markup["x2"], markup["y2"]],
                        fill=markup["color"],
                        width=2
                    )
        
        # Get comment if enabled
        comment = ""
        if self.add_comment_var.get():
            comment = self.get_comment()
            if comment is None:
                return
        
        # Store screenshot with markups applied
        timestamp = datetime.now()
        self.screenshots.append({
            'image': marked_image,
            'comment': comment,
            'timestamp': timestamp
        })
        
    def get_comment(self):
        """Get a comment from the user for the screenshot"""
        comment = simpledialog.askstring("Add comment for the screenshot", 
                                       "Enter a comment for this screenshot:",
                                       parent=self.root)
        return comment if comment is not None else ""
    
    def show_help(self):
        """Show help information dialog"""
        help_text = """SnipIT v2.0 - Support & Contact

KEYBOARD SHORTCUTS:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Ctrl+Alt+F  - Full Screen Capture
Ctrl+Alt+P  - Partial Capture (with selection)

Perfect for capturing expanded dropdowns without
losing focus to the web application!

For Support and Queries please reach out to:

Work Email:
nanthish.t@gds.ey.com

Personal Email:
nanthishwaran579@gmail.com

"""
        messagebox.showinfo("SnipIT Help", help_text)
    
    def close_tool(self):
        """Close the tool with confirmation"""
        if messagebox.askyesno("Close", "Are you sure you want to close SnipIT?"):
            self.root.quit()
        
    def end_session(self):
        """End the session and create Word document"""
        if not self.screenshots:
            messagebox.showinfo("No Screenshots", "No screenshots were taken.")
            return
            
        try:
            # Create Word document
            doc = Document()
            
            # Add screenshots with comments and timestamps
            for i, shot in enumerate(self.screenshots, 1):
                # Add screenshot number as text with blue color and bold
                p = doc.add_paragraph()
                run = p.add_run(f"Screenshot {i}")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                
                # Add timestamp
                timestamp_text = shot["timestamp"].strftime("%Y-%m-%d %H:%M:%S")
                doc.add_paragraph(f"Timestamp: {timestamp_text}")
                
                # Add comment only if it exists (no blank comment section)
                if shot["comment"]:
                    doc.add_paragraph(shot["comment"])
                
                # Save image to temporary file and add to document
                temp_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                shot["image"].save(temp_file.name, "PNG")
                temp_file.close()
                
                # Add image to document
                doc.add_picture(temp_file.name, width=Inches(5.5))
                
                # Clean up temp file
                os.unlink(temp_file.name)
                
                # Add page break except for last screenshot
                if i < len(self.screenshots):
                    doc.add_page_break()
            
            # Save to a temporary file
            temp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_docx_path = temp_docx.name
            temp_docx.close()
            doc.save(temp_docx_path)
            
            # Open in Word using COM interface
            try:
                word = win32com.client.GetObject(Class="Word.Application")
            except:
                # If Word isn't already running, start it
                word = win32com.client.Dispatch("Word.Application")
            
            word.Visible = True
            
            # Open the document
            word_doc = word.Documents.Open(FileName=os.path.abspath(temp_docx_path))
            
            # Bring Word to foreground
            word.Activate()
            
            # Clean up - delete temp file after a delay to ensure Word has loaded it
            def cleanup_temp():
                time.sleep(2)
                try:
                    os.unlink(temp_docx_path)
                except:
                    pass
            
            cleanup_thread = threading.Thread(target=cleanup_temp, daemon=True)
            cleanup_thread.start()
            
            # Automatically quit the tool after a short delay
            self.root.after(500, self.root.quit)
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create document: {str(e)}")
            
    def run(self):
        """Start the application"""
        self.root.mainloop()


def main():
    """Main entry point"""
    app = ScreenshotTool()
    app.run()


if __name__ == "__main__":
    main()
